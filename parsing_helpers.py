# sections/parsing_helpers.py
import re
import pandas as pd
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

# -------- REGEX patterns --------
EPIC_RE = re.compile(r"^\s*Epic\s+(\d+)\s*[:\-\u2013\u2014]\s*(.+)\s*$", re.IGNORECASE)
STORY_RE = re.compile(r"^\s*(?:User\s+)?Story\s+(\d+(?:\.\d+)*)\s*[:\-\u2013\u2014]\s*(.+)\s*$", re.IGNORECASE)

HEADER_ALIASES = {
    "scenario": "Scenario", "given": "Given", "precondition": "Given",
    "when": "When", "action": "When", "then": "Then",
    "expected": "Expected", "expected result": "Expected", "result": "Expected",
    "acceptance criteria": "Acceptance Criteria", "criteria": "Acceptance Criteria", "ac": "Acceptance Criteria",
    "#": "AC #", "no": "AC #", "id": "AC #", "sr no": "AC #", "s no": "AC #",
    "sno": "AC #", "srno": "AC #", "ac #": "AC #", "ac no": "AC #", "ac number": "AC #",
}

# -------- Helpers --------
def iter_block_items(parent):
    """Yield paragraphs and tables from a Word document in order."""
    xml_children = list(parent.element.body.iterchildren())
    table_idx = 0
    for idx, child in enumerate(xml_children):
        if isinstance(child, CT_P):
            p_count = sum(1 for c in xml_children[: idx + 1] if isinstance(c, CT_P))
            yield ("p", parent.paragraphs[p_count - 1])
        elif isinstance(child, CT_Tbl):
            if table_idx < len(parent.tables):
                yield ("t", parent.tables[table_idx])
            table_idx += 1

def _canon_header(text: str) -> str:
    raw = (text or "").strip()
    t = raw.lower()
    t = re.sub(r"[^\w#]+", " ", t)
    t = re.sub(r"\s+", " ", t).strip()
    return HEADER_ALIASES.get(t, raw if raw else "")

def looks_like_ac_table(table):
    keywords = {"acceptance", "criteria", "scenario", "given", "when", "then", "expected", "result"}
    for hdr_idx in (0, 1):
        if hdr_idx >= len(table.rows):
            break
        cells_text = " | ".join(cell.text.strip().lower() for cell in table.rows[hdr_idx].cells)
        if any(k in cells_text for k in keywords):
            return True, hdr_idx
    return False, None

def _row_is_empty(row) -> bool:
    return all(not (cell.text or "").strip() for cell in row.cells)

def _first_data_row_index(table, header_row_index):
    total_rows = len(table.rows)
    data_start = header_row_index + 1
    if data_start < total_rows and _row_is_empty(table.rows[data_start]):
        data_start += 1
    return min(data_start, total_rows)

def count_ac_rows(table, header_row_index):
    total_rows = len(table.rows)
    data_start = _first_data_row_index(table, header_row_index)
    count = sum(1 for r in range(data_start, total_rows) if not _row_is_empty(table.rows[r]))
    return max(count, 0)

def parse_ac_table_rows_minimal(table, header_row_index):
    header_cells = table.rows[header_row_index].cells
    headers = [_canon_header(c.text) for c in header_cells]

    def _find_idx(names):
        for i, h in enumerate(headers):
            if h in names: return i
        return None

    idx_acnum = _find_idx({"AC #"})
    idx_scenario = _find_idx({"Scenario"})
    idx_free_ac = _find_idx({"Acceptance Criteria"})

    data_start = _first_data_row_index(table, header_row_index)
    out = []
    for r_idx in range(data_start, len(table.rows)):
        row = table.rows[r_idx]
        if _row_is_empty(row): continue
        cells = [cell.text.strip() for cell in row.cells]
        ac_no = cells[idx_acnum].strip() if idx_acnum is not None and idx_acnum < len(cells) else ""
        scenario = cells[idx_scenario].strip() if idx_scenario is not None and idx_scenario < len(cells) else ""
        if not scenario and idx_free_ac is not None and idx_free_ac < len(cells):
            scenario = cells[idx_free_ac].strip()
        if ac_no or scenario:
            out.append({"AC #": ac_no, "Scenario": scenario})
    return out

def extract_user_stories_and_acs(docx_file):
    doc = Document(docx_file)
    paragraphs = doc.paragraphs
    if len(paragraphs) == 0:
        return pd.DataFrame(columns=["Module","Epic","Story ID","Story Title","Acceptance Criteria Count"]), \
               pd.DataFrame(columns=["Module","Epic","Story ID","Story Title","AC #","Scenario"])

    paragraphs_nonempty = [p.text.strip() for p in paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs_nonempty)
    module_match = re.search(r"Module\s*[:\-\u2013\u2014]\s*(.+)", full_text, flags=re.IGNORECASE)
    module = module_match.group(1).strip() if module_match else "Unknown"

    stories, ac_rows = [], []
    current_epic, current_story = None, None

    for kind, obj in iter_block_items(doc):
        if kind == "p":
            line = (obj.text or "").strip()
            if not line: continue
            em = EPIC_RE.match(line)
            if em:
                current_epic = f"{em.group(1)}: {em.group(2).strip()}"
                continue
            sm = STORY_RE.match(line)
            if sm:
                current_story = {"Module": module, "Epic": current_epic or "Unknown",
                                 "Story ID": sm.group(1).strip(), "Story Title": sm.group(2).strip(),
                                 "Acceptance Criteria Count": 0}
                stories.append(current_story)
        elif kind == "t" and current_story:
            is_ac, hdr_idx = looks_like_ac_table(obj)
            if not is_ac: continue
            current_story["Acceptance Criteria Count"] += count_ac_rows(obj, hdr_idx)
            for entry in parse_ac_table_rows_minimal(obj, hdr_idx):
                ac_rows.append({"Module": current_story["Module"], "Epic": current_story["Epic"],
                                "Story ID": current_story["Story ID"], "Story Title": current_story["Story Title"],
                                **entry})
    return pd.DataFrame(stories), pd.DataFrame(ac_rows)
